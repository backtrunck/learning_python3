# -*- coding: utf-8 -*-
"""
Created on Wed Mar 14 14:33:06 2018

@author: luiscar
"""
import datetime
from docx import Document

def explorar_table(doc):
    for index,t in enumerate(doc.tables):
        print('Tabela[{}]'.format(index))
        for i in range(len(t.rows)):
            for j in range(len(t.columns)):
                print('Celula[{},{}]'.format(i,j))
                try:
                    if t.cell(i,j).tables:
                        explorar_table(t.cell(i,j))
                    else:
                        print(t.cell(i,j).text)
                except IndexError as e:
                    print(e)
                    


def captura_dados(tabela):
    
    #if tabela.cell(0,1).tables[0].cell(0,0).text.strip() != "NOME  MATRÍCULA":
    #    return {}
    if not tabela.cell(0,1).tables:
        return {}
    ficha = {}
    valor = tabela.cell(0,1).tables[0].cell(0,1).text
    
    ficha["nome"] = repr(valor.strip())
    valor = tabela.cell(1,1).tables[0].cell(0,1).text
    
    ficha["pai_mae"] = repr(valor.strip())
    
    valor = tabela.cell(2,1).tables[0].cell(0,0).text
    
    ficha["cpf_outros"] = repr(valor.strip())
    
    valor = tabela.cell(2,1).tables[0].cell(0,2).text
    
    ficha["serie_zona"] = repr(valor.strip())
    
    
    #emissão,seção,cadastro
    #valor = tabela.cell(2,1).tables[0].cell(0,3).text
    #ficha["emissao_outros"] = repr(valor.strip())
  
    
    #admissão, opção fgts, cargo, seriviços
    valor = tabela.cell(3,0).text
    ficha["admissao_outros"] = repr(valor.strip())
    
    
    #forma de pagamento, jornada, seção salario
    valor = tabela.cell(3,6).text
    ficha["forma_pagamento_outros"] = repr(valor.strip())
       
    
    #data de nascimento, nacionalidade
    valor = tabela.cell(4,0).text
    ficha["data_nascimento_outros"] = repr(valor.strip())
    
    #Estado civil, naturalidade
    valor = tabela.cell(4,3).text
    ficha["estado_civil_outros"] = repr(valor.strip())
    
    #Sexo, grau de intrução, estado natal
    valor = tabela.cell(4,6 ).text
    ficha["grau_instrucao_outros"] = repr(valor.strip())
    
    #Quando estrangeiro,data chegada, tipo visto
    valor = tabela.cell(5,0).text
    ficha["qdo_estrangeiro_outros"] = repr(valor.strip())
        
    #conjuge brasileiro, registro geral,validade carteira
    valor = tabela.cell(5,3).text
    ficha["conjuge_outros"] = repr(valor.strip())
        
    #identidade, numero decreto, validade 
    valor = tabela.cell(5,9).text
    ficha["decreto_outros"] = repr(valor.strip())
        
    i=7
    
    ficha["dependente"] = []
    
    while tabela.cell(i,0).text.strip() != 'ENDEREÇO':
        dependente = {}
        valor = tabela.cell(i,0).text
        dependente["nome"] = valor.strip()
        valor = tabela.cell(i,9).text
        dependente["estado_civil"] = valor.strip()
        valor = tabela.cell(i,10).text
        dependente["parentesco"] = valor.strip()
        ficha["dependente"].append(dependente)
        i += 1
             
    i += 1
    
    #endereço
    valor = tabela.cell(i,0).text
    ficha["Endereco"] = valor.strip()
    
    #numero
    valor = tabela.cell(i,5).text
    ficha["Endereco_numero"] = valor.strip()
    
    
    #bairro
    valor = tabela.cell(i,7).text
    ficha["Endereco_bairro"] = valor.strip()
                      
    #cidade
    valor = tabela.cell(i,10).text
    ficha["Endereco_cidade"] = valor.strip()
    
    i += 1
    
    valor = tabela.cell(i,0).text.strip()
    
    i += 1
    
    if valor[:6] == 'FÉRIAS':
        ficha["ferias"]=[]
        
        while tabela.cell(i,0).text.strip() != 'ALTERAÇÕES DE SALÁRIO':
            periodo_ferias={}
            valor = tabela.cell(i,0).text
            periodo_ferias["periodo_aquisitivo"] = valor.strip()
            valor = tabela.cell(i,5).text
            periodo_ferias["periodo_gozo"] = valor.strip()
            ficha["ferias"].append(periodo_ferias)
            i += 1
            
    valor = tabela.cell(i,0).text.strip()
    
    if valor == 'ALTERAÇÕES DE SALÁRIO':
        i += 1
        
        ficha["alteracoes_salarios"] = []
        
        while tabela.cell(i,0).text.strip() == '':
            alteracao_salario = {}
            valor = tabela.cell(i,2).text
            alteracao_salario["data"] = valor.strip()   
            valor = tabela.cell(i,6).text
            alteracao_salario["salario"] = valor.strip()   
            valor = tabela.cell(i,10).text
            alteracao_salario["motivo"] = valor.strip()
            ficha["alteracoes_salarios"].append(alteracao_salario)
            i += 1
    
    valor = tabela.cell(i,0).text.strip()
        
    if valor == "CONTRIBUIÇÃO SINDICAL":
        i += 1
        ficha["tipo_contribuicao_sindical"] = tabela.cell(i,0).\
                                                 text.strip()
        ficha["data_contribuicao_sindical"] = tabela.cell(i,2).\
                                                 text.strip()
        ficha["valor_contribuicao_sindical"] = tabela.cell(i,10).\
                                                 text.strip()
        i += 1
    
    valor = tabela.cell(i,0).text.strip() 
    
   
         
         
    if valor == "ALTERAÇÕES DE FUNÇÃO":
        i += 1
        ficha["alteracoes_funcao"] = []
        
        while tabela.cell(i,0).text.strip() == '':
            alteracao_funcao = {}
            valor = tabela.cell(i,2).text
            alteracao_funcao["data"] = valor.strip()
            valor = tabela.cell(i,5).text
            alteracao_funcao["funcao"] = valor.strip()
            valor = tabela.cell(i,10).text
            alteracao_funcao["motivo"] = valor.strip()
            ficha["alteracoes_funcao"].append(alteracao_funcao)
            i += 1
            
    valor = tabela.cell(i,0).text.strip() 
    
    if valor == "MUDANÇAS DE SEÇÃO": 
        
        i += 1
        
        ficha["mudancas_secao"] = []
        
        while tabela.cell(i,0).text.strip() == '':
            mudanca_secao = {}
            valor = tabela.cell(i,2).text
            mudanca_secao["data"] = valor.strip()  
            valor = tabela.cell(i,5).text
            mudanca_secao["secao"] = valor.strip()
            valor = tabela.cell(i,10).text
            mudanca_secao["motivo"] = valor.strip()
            ficha["mudancas_secao"].append(mudanca_secao)
            i += 1
            
    valor = tabela.cell(i,0).text.strip() 
    
    if valor == "HORÁRIO ATUAL": 
        
        i += 1
        ficha["horario_atual"] = []
        while  tabela.cell(i,0).text.strip() == '':
            horario = {}
            horario["data_horario_atual"] = tabela.cell(i,2).text.strip() 
            horario["desc_horario_atual"] = tabela.cell(i,5).text.strip() 
            ficha["horario_atual"].append(horario)
            i += 1
   
    valor = tabela.cell(i,0).text.strip() 
    
    ficha["teste2"] = tabela.cell(i,0).text.strip() 
    
    if valor == "ANOTAÇÕES GERAIS": 
        
        ficha["desc_anotacoes_gerais"] = tabela.cell(i,0).text.strip()
        ficha["data_anotacoes_gerais"] = tabela.cell(i,2).text.strip()
        
        i += 1
    
    ficha["data_demissao_outros"] = tabela.cell(i,0).text.strip()

    
    return ficha    
                 
if __name__ == '__main__':
    dc = Document('C:\\Users\\luiscar\\Documents\\profissional\\cgu-ba\\' + \
                       'nae\\operacao_melinoe\\melinoe\\ficha.registro.' + \
                       'demitidos.reduzido.docx')
    
    sufixo_data = datetime.datetime.now().strftime("%Y_%m_%d")
    f = open("teste_docx_%s.%s" % (sufixo_data,".txt"),"w")
    
    for t in range(len(dc.tables) - 1):
        print(len(dc.tables))
        ficha = captura_dados(dc.tables[t])
        
        for u,v in ficha.items():
            if v is list:
                for i in v:
                    if i is dict:
                        for ch,vl in i.items():
                            f.writelines("%s : %s" % (ch,vl))
                            f.writelines("\n")
                else:
                    f.writelines("%s : %s" % (u,v))
                    f.writelines("\n")
    f.close()
    
    for i,j in ficha.items():
        
        print(i,":",j)